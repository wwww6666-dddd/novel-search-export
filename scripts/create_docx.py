"""Create Word doc with clickable TOC and content-type formatting.
Usage: python create_docx.py <input_json> <output_docx> [--title TITLE] [--author AUTHOR]
"""
import sys, json, re
from html import unescape
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Microsoft YaHei"

def set_font(run, size=Pt(12), bold=False, color=None, italic=False):
    run.font.name = FONT
    run.font.size = size
    run.bold = bold
    run.italic = italic
    run.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if color: run.font.color.rgb = color

def add_bookmark(paragraph, name):
    bid = str(abs(hash(name)) % (2**31 - 1))
    bs = OxmlElement("w:bookmarkStart")
    bs.set(qn("w:id"), bid); bs.set(qn("w:name"), name)
    be = OxmlElement("w:bookmarkEnd")
    be.set(qn("w:id"), bid); be.set(qn("w:name"), name)
    run = paragraph.add_run()
    run._r.addprevious(bs); run._r.addnext(be)

def add_hyperlink(paragraph, text, bookmark):
    rId = paragraph.part.relate_to(paragraph.part,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=False)
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("w:anchor"), bookmark)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    c = OxmlElement("w:color"); c.set(qn("w:val"), "2E75B6"); rpr.append(c)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rpr.append(u)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "24"); rpr.append(sz)
    rf = OxmlElement("w:rFonts"); rf.set(qn("w:eastAsia"), FONT); rpr.append(rf)
    run.append(rpr)
    run.text = text
    hl.append(run)
    paragraph._p.append(hl)

def split_sentences(text):
    parts = re.split(r"(?<=[。！？])s*", text)
    return [p.strip() for p in parts if p.strip()] or [text.strip()]

def add_para(doc, text):
    if not text.strip(): return
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)

    t = text.strip()
    if t.startswith("【") or t.startswith("［"):
        p.paragraph_format.left_indent = Cm(0.5)
        r = p.add_run(t)
        set_font(r, size=Pt(12), bold=True, color=RGBColor(0x00, 0x70, 0xC0))
    elif re.search(r'[""][^""]*[""]', t):
        segs = re.split(r'(["][^"]*["])', t)
        for seg in segs:
            if not seg: continue
            r = p.add_run(seg)
            set_font(r, size=Pt(12))
            if (seg.startswith(""") or seg.startswith(""")) and len(seg) < 80:
                r.font.color.rgb = RGBColor(0x33, 0x33, 0x99)
    else:
        r = p.add_run(t)
        set_font(r, size=Pt(12))

def sort_chapters(chapters):
    """Sort chapters: numerically by chapter number first, then extras at end."""
    main, extra = [], []
    for ch in chapters:
        title = ch.get("title", "")
        m = re.search(r"第s*(d+)s*章", title)
        if m:
            main.append((int(m.group(1)), ch))
        else:
            extra.append(ch)
    main.sort(key=lambda x: x[0])
    result = [c for _, c in main]
    result.extend(extra)
    return result

def create_docx(chapters, output_path, book_title="", author=""):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT; style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    for sec in doc.sections:
        sec.top_margin = Cm(2); sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)

    # Sort chapters properly
    chapters = sort_chapters(chapters)

    # == Cover ==
    for _ in range(6): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(book_title); set_font(r, size=Pt(22), bold=True)
    if author:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        r = p.add_run(author); set_font(r, size=Pt(14), color=RGBColor(0x55,0x55,0x55))

    # == TOC ==
    doc.add_page_break()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("目  录"); set_font(r, size=Pt(18), bold=True)
    p.paragraph_format.space_after = Pt(18)

    # Detect chapter groups by common patterns
    main_chapters = []
    fanwai_groups = {}  # {group_name: [chapters]}

    for ch in chapters:
        t = ch.get("title", "")
        # Check for group markers in title
        m_group = re.search(r"[(（]([^)）]+番外[^)）]*)[)）]", t)
        if m_group:
            gname = m_group.group(1)
            fanwai_groups.setdefault(gname, []).append(ch)
            continue
        # Simple extra chapter patterns
        if any(kw in t for kw in ["番外", "后记", "尾声", "结语", "完结"]):
            fanwai_groups.setdefault("番外/后记", []).append(ch)
        elif re.search(r"第s*d+s*章", t):
            main_chapters.append(ch)
        else:
            fanwai_groups.setdefault("其他", []).append(ch)

    # Output TOC: main chapters first, then each group
    for i, ch in enumerate(main_chapters):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        t = ch["title"] if len(ch["title"]) <= 55 else ch["title"][:52] + "..."
        add_hyperlink(p, t, f"ch_{i}")

    # Add group headers and entries
    base_i = len(main_chapters)
    for gname, gchapters in fanwai_groups.items():
        if not gchapters: continue
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"--- {gname} ---")
        set_font(r, size=Pt(12), bold=True, color=RGBColor(0x33,0x33,0x33))
        for j, ch in enumerate(gchapters):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Cm(0.5)
            t = ch["title"] if len(ch["title"]) <= 55 else ch["title"][:52] + "..."
            add_hyperlink(p, t, f"ch_{base_i + j}")
        base_i += len(gchapters)

    # == Chapter content ==
    doc.add_page_break()
    for i, ch in enumerate(chapters):
        title = ch.get("title", "")
        content = unescape(ch.get("content", ""))
        content = content.replace("r", "")

        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title); set_font(r, size=Pt(16), bold=True)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(14)
        add_bookmark(p, f"ch_{i}")

        for sent in split_sentences(content):
            add_para(doc, sent)

        sep = doc.add_paragraph(); sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = sep.add_run("─" * 7)
        set_font(r, size=Pt(10), color=RGBColor(0xBB, 0xBB, 0xBB))
        doc.add_page_break()

    doc.save(output_path)
    total = sum(len(unescape(ch.get("content", ""))) for ch in chapters)
    print(f"Saved: {output_path}")
    print(f"  Chapters: {len(chapters)}, Total chars: {total:,}")

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python create_docx.py <input_json> <output_docx> [--title TITLE] [--author AUTHOR]")
        sys.exit(1)
    with open(sys.argv[1], "r", encoding="utf-8") as f:
        chapters = json.load(f)
    output = sys.argv[2]
    title = sys.argv[sys.argv.index("--title") + 1] if "--title" in sys.argv else ""
    author = sys.argv[sys.argv.index("--author") + 1] if "--author" in sys.argv else ""
    create_docx(chapters, output, title, author)

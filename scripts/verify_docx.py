"""verify_docx.py - Export pre-check: cover, TOC links, chapter headers, encoding"""
import re, sys
from docx import Document

def verify(path):
    d = Document(path)
    paras = [p for p in d.paragraphs if p.text.strip()]
    if not paras:
        print("FAIL: Empty document"); return False

    # 1. Cover title
    print(f"Cover: {paras[0].text[:40]}")

    # 2. TOC hyperlinks
    toc = sum(1 for p in d.paragraphs for r in p._p if r.tag.endswith("hyperlink"))
    if toc < 5:
        print(f"WARN: Only {toc} TOC links"); return False
    print(f"TOC links: {toc}")

    # 3. Chapter headers (accept multiple patterns)
    chs = [p for p in d.paragraphs
           if (re.match(r"第\s*\d+\s*章", p.text.strip())
               or any(kw in p.text for kw in ["番外", "后记", "尾声"]))
           and p.alignment == 1]
    print(f"Chapter headers: {len(chs)}")

    # 4. Content volume
    text = "\n".join(p.text for p in d.paragraphs)
    if len(text) < 10000:
        print("FAIL: Too little text"); return False
    print(f"Total chars: {len(text)}")

    # 5. Encoding
    try: text.encode("utf-8")
    except: print("FAIL: Encoding error"); return False
    print("Encoding: OK")

    print(f"\nResult: OK (file size: {len(text)} chars)")
    return True

if __name__ == "__main__":
    path = sys.argv[1] if len(sys.argv) > 1 else ""
    if not path:
        import glob
        candidates = glob.glob("D:/小说/小说/*.docx")
        path = candidates[0] if candidates else ""
    sys.exit(0 if verify(path) else 1)
